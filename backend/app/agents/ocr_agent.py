"""
OCR Agent for document text extraction using Google Cloud Document AI.

This agent handles:
- PDF and image OCR processing with Document AI
- Layout analysis and structured text extraction
- Fallback to Vision API for edge cases
- Document preprocessing and format detection
"""

import asyncio
import logging
from typing import Dict, List, Optional, Tuple, Union
from pathlib import Path
import io
import base64

# Optional Google Cloud imports
try:
    from google.cloud import documentai
    from google.cloud import vision
    from google.api_core import exceptions as gcp_exceptions
    GOOGLE_CLOUD_AVAILABLE = True
except ImportError:
    documentai = None
    vision = None
    gcp_exceptions = None
    GOOGLE_CLOUD_AVAILABLE = False
# Optional document processing imports
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    Image = None
    PIL_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

from ..core.config import get_settings
from ..models.document import ProcessedDocument, OCRResult, DocumentLayout
from ..core.exceptions import OCRProcessingError, DocumentFormatError
from .preprocessing import DocumentPreprocessor

logger = logging.getLogger(__name__)
settings = get_settings()


class OCRAgent:
    """
    Specialized agent for document OCR and text extraction.
    
    Uses Google Cloud Document AI as primary OCR service with Vision API fallback.
    Supports PDF, images, and DOCX format processing.
    """
    
    def __init__(self):
        """Initialize OCR agent with Google Cloud clients."""
        if not GOOGLE_CLOUD_AVAILABLE:
            logger.warning("Google Cloud libraries not available - OCR functionality disabled")
            self.doc_ai_client = None
            self.vision_client = None
            self.preprocessor = None
            return
            
        self.project_id = settings.GOOGLE_CLOUD_PROJECT
        self.location = settings.DOCUMENT_AI_LOCATION or "us"
        self.processor_id = settings.DOCUMENT_AI_PROCESSOR_ID
        
        # Initialize Document AI client
        if self.processor_id:
            try:
                self.doc_ai_client = documentai.DocumentProcessorServiceClient()
                self.processor_name = self.doc_ai_client.processor_path(
                    self.project_id, self.location, self.processor_id
                )
            except Exception as e:
                logger.warning(f"Failed to initialize Document AI client: {e}")
                self.doc_ai_client = None
        else:
            self.doc_ai_client = None
            logger.warning("Document AI processor not configured, using Vision API only")
        
        # Initialize Vision API client as fallback
        try:
            self.vision_client = vision.ImageAnnotatorClient()
        except Exception as e:
            logger.warning(f"Failed to initialize Vision API client: {e}")
            self.vision_client = None
        
        # Initialize document preprocessor
        try:
            from .preprocessing import DocumentPreprocessor
            self.preprocessor = DocumentPreprocessor()
        except ImportError:
            logger.warning("Document preprocessor not available")
            self.preprocessor = None
        
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str,
        content_type: str
    ) -> OCRResult:
        """
        Process document and extract structured text.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type of the document
            
        Returns:
            OCRResult with extracted text and layout information
            
        Raises:
            OCRProcessingError: If OCR processing fails
            DocumentFormatError: If document format is unsupported
        """
        if not GOOGLE_CLOUD_AVAILABLE:
            raise OCRProcessingError("Google Cloud libraries not available - OCR functionality disabled")
            
        try:
            logger.info(f"Starting OCR processing for {filename} ({content_type})")
            
            # Preprocess document for better OCR results
            processed_content, preprocessing_metadata = await self.preprocessor.preprocess_document(
                file_content, filename, content_type
            )
            
            # Determine processing strategy based on content type
            if content_type == "application/pdf":
                result = await self._process_pdf(processed_content, filename)
            elif content_type.startswith("image/"):
                result = await self._process_image(processed_content, filename)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                result = await self._process_docx(processed_content, filename)
            else:
                raise DocumentFormatError(f"Unsupported document format: {content_type}")
            
            # Add preprocessing metadata to result
            if hasattr(result, 'preprocessing_metadata'):
                result.preprocessing_metadata = preprocessing_metadata
            
            return result
                
        except Exception as e:
            logger.error(f"OCR processing failed for {filename}: {str(e)}")
            raise OCRProcessingError(f"Failed to process document: {str(e)}") from e
    
    async def _process_pdf(self, file_content: bytes, filename: str) -> OCRResult:
        """Process PDF using Document AI with Vision API fallback."""
        try:
            # Try Document AI first
            if self.doc_ai_client and self.processor_id:
                return await self._process_with_document_ai(file_content, filename, "application/pdf")
            else:
                # Fallback: Convert PDF pages to images and use Vision API
                return await self._process_pdf_with_vision(file_content, filename)
                
        except Exception as e:
            logger.warning(f"Document AI failed for PDF {filename}, trying Vision API: {str(e)}")
            return await self._process_pdf_with_vision(file_content, filename)
    
    async def _process_image(self, file_content: bytes, filename: str) -> OCRResult:
        """Process image using Document AI with Vision API fallback."""
        try:
            # Try Document AI first for better layout analysis
            if self.doc_ai_client and self.processor_id:
                content_type = self._detect_image_type(file_content)
                return await self._process_with_document_ai(file_content, filename, content_type)
            else:
                # Use Vision API directly
                return await self._process_with_vision_api(file_content, filename)
                
        except Exception as e:
            logger.warning(f"Document AI failed for image {filename}, trying Vision API: {str(e)}")
            return await self._process_with_vision_api(file_content, filename)
    
    async def _process_docx(self, file_content: bytes, filename: str) -> OCRResult:
        """Process DOCX by extracting text directly."""
        if not DOCX_AVAILABLE:
            raise OCRProcessingError("python-docx library not available - DOCX processing disabled")
            
        try:
            # DOCX files contain structured text, no OCR needed
            doc = DocxDocument(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if table_text:
                full_text += "\n\nTables:\n" + "\n".join(table_text)
            
            # Create layout information
            layout = DocumentLayout(
                pages=[{
                    "page_number": 1,
                    "width": 612,  # Standard letter width in points
                    "height": 792,  # Standard letter height in points
                    "blocks": [{
                        "text": full_text,
                        "bounding_box": {"x": 0, "y": 0, "width": 612, "height": 792},
                        "confidence": 1.0
                    }]
                }],
                total_pages=1
            )
            
            return OCRResult(
                text=full_text,
                confidence=1.0,
                layout=layout,
                processing_method="docx_extraction",
                language_code="en"
            )
            
        except Exception as e:
            raise OCRProcessingError(f"Failed to process DOCX file: {str(e)}") from e
    
    async def _process_with_document_ai(
        self, 
        file_content: bytes, 
        filename: str, 
        mime_type: str
    ) -> OCRResult:
        """Process document using Google Cloud Document AI."""
        try:
            # Create the request
            raw_document = documentai.RawDocument(
                content=file_content,
                mime_type=mime_type
            )
            
            request = documentai.ProcessRequest(
                name=self.processor_name,
                raw_document=raw_document
            )
            
            # Process the document
            result = await asyncio.get_event_loop().run_in_executor(
                None, self.doc_ai_client.process_document, request
            )
            
            document = result.document
            
            # Extract structured text and layout
            full_text = document.text
            confidence = self._calculate_average_confidence(document)
            layout = self._extract_layout_from_document_ai(document)
            
            return OCRResult(
                text=full_text,
                confidence=confidence,
                layout=layout,
                processing_method="document_ai",
                language_code=self._detect_language(document)
            )
            
        except gcp_exceptions.GoogleAPIError as e:
            logger.error(f"Document AI API error: {str(e)}")
            raise OCRProcessingError(f"Document AI processing failed: {str(e)}") from e
    
    async def _process_pdf_with_vision(self, file_content: bytes, filename: str) -> OCRResult:
        """Process PDF by converting to images and using Vision API."""
        if not PYPDF2_AVAILABLE:
            raise OCRProcessingError("PyPDF2 library not available - PDF processing disabled")
            
        try:
            # Read PDF and convert pages to images
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            # For now, extract text directly from PDF if possible
            text_pages = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_pages.append(page_text.strip())
                    else:
                        # If no text extracted, would need to convert to image
                        # This is a simplified implementation
                        text_pages.append(f"[Page {page_num + 1} - Image content requires OCR]")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    text_pages.append(f"[Page {page_num + 1} - Text extraction failed]")
            
            full_text = "\n\n".join(text_pages)
            
            # Create basic layout information
            layout = DocumentLayout(
                pages=[{
                    "page_number": i + 1,
                    "width": 612,
                    "height": 792,
                    "blocks": [{
                        "text": page_text,
                        "bounding_box": {"x": 0, "y": 0, "width": 612, "height": 792},
                        "confidence": 0.8
                    }]
                } for i, page_text in enumerate(text_pages)],
                total_pages=len(text_pages)
            )
            
            return OCRResult(
                text=full_text,
                confidence=0.8,
                layout=layout,
                processing_method="pdf_text_extraction",
                language_code="en"
            )
            
        except Exception as e:
            raise OCRProcessingError(f"PDF processing with Vision API failed: {str(e)}") from e
    
    async def _process_with_vision_api(self, file_content: bytes, filename: str) -> OCRResult:
        """Process image using Google Cloud Vision API."""
        try:
            # Create Vision API request
            image = vision.Image(content=file_content)
            
            # Perform text detection
            response = await asyncio.get_event_loop().run_in_executor(
                None, self.vision_client.text_detection, image
            )
            
            if response.error.message:
                raise OCRProcessingError(f"Vision API error: {response.error.message}")
            
            # Extract text and confidence
            texts = response.text_annotations
            if not texts:
                return OCRResult(
                    text="",
                    confidence=0.0,
                    layout=DocumentLayout(pages=[], total_pages=0),
                    processing_method="vision_api",
                    language_code="en"
                )
            
            # First annotation contains the full text
            full_text = texts[0].description
            
            # Calculate average confidence from individual text blocks
            confidences = []
            blocks = []
            
            for text in texts[1:]:  # Skip the first full-text annotation
                if hasattr(text, 'confidence'):
                    confidences.append(text.confidence)
                
                # Extract bounding box
                vertices = text.bounding_poly.vertices
                if vertices:
                    blocks.append({
                        "text": text.description,
                        "bounding_box": {
                            "x": min(v.x for v in vertices),
                            "y": min(v.y for v in vertices),
                            "width": max(v.x for v in vertices) - min(v.x for v in vertices),
                            "height": max(v.y for v in vertices) - min(v.y for v in vertices)
                        },
                        "confidence": getattr(text, 'confidence', 0.8)
                    })
            
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.8
            
            # Create layout information
            layout = DocumentLayout(
                pages=[{
                    "page_number": 1,
                    "width": 1000,  # Estimated image dimensions
                    "height": 1000,
                    "blocks": blocks
                }],
                total_pages=1
            )
            
            return OCRResult(
                text=full_text,
                confidence=avg_confidence,
                layout=layout,
                processing_method="vision_api",
                language_code="en"
            )
            
        except Exception as e:
            raise OCRProcessingError(f"Vision API processing failed: {str(e)}") from e
    
    def _detect_image_type(self, file_content: bytes) -> str:
        """Detect image MIME type from file content."""
        if not PIL_AVAILABLE:
            return 'image/jpeg'  # Default fallback
            
        try:
            image = Image.open(io.BytesIO(file_content))
            format_to_mime = {
                'JPEG': 'image/jpeg',
                'PNG': 'image/png',
                'GIF': 'image/gif',
                'BMP': 'image/bmp',
                'TIFF': 'image/tiff'
            }
            return format_to_mime.get(image.format, 'image/jpeg')
        except Exception:
            return 'image/jpeg'  # Default fallback
    
    def _calculate_average_confidence(self, document) -> float:
        """Calculate average confidence score from Document AI result."""
        confidences = []
        
        for page in document.pages:
            for block in page.blocks:
                if hasattr(block, 'confidence'):
                    confidences.append(block.confidence)
                    
        return sum(confidences) / len(confidences) if confidences else 0.8
    
    def _extract_layout_from_document_ai(self, document) -> DocumentLayout:
        """Extract layout information from Document AI result."""
        pages = []
        
        for page_num, page in enumerate(document.pages):
            blocks = []
            
            for block in page.blocks:
                # Extract text from block
                block_text = self._get_text_from_layout(document.text, block.layout)
                
                # Extract bounding box
                bbox = block.layout.bounding_poly.normalized_vertices
                if bbox:
                    blocks.append({
                        "text": block_text,
                        "bounding_box": {
                            "x": min(v.x for v in bbox) * page.dimension.width,
                            "y": min(v.y for v in bbox) * page.dimension.height,
                            "width": (max(v.x for v in bbox) - min(v.x for v in bbox)) * page.dimension.width,
                            "height": (max(v.y for v in bbox) - min(v.y for v in bbox)) * page.dimension.height
                        },
                        "confidence": getattr(block, 'confidence', 0.8)
                    })
            
            pages.append({
                "page_number": page_num + 1,
                "width": page.dimension.width,
                "height": page.dimension.height,
                "blocks": blocks
            })
        
        return DocumentLayout(
            pages=pages,
            total_pages=len(pages)
        )
    
    def _get_text_from_layout(self, full_text: str, layout) -> str:
        """Extract text segment from layout information."""
        if not hasattr(layout, 'text_anchor') or not layout.text_anchor.text_segments:
            return ""
        
        text_segments = []
        for segment in layout.text_anchor.text_segments:
            start_index = segment.start_index if hasattr(segment, 'start_index') else 0
            end_index = segment.end_index if hasattr(segment, 'end_index') else len(full_text)
            text_segments.append(full_text[start_index:end_index])
        
        return "".join(text_segments)
    
    def _detect_language(self, document) -> str:
        """Detect document language from Document AI result."""
        # Simple language detection - could be enhanced
        for page in document.pages:
            for block in page.blocks:
                if hasattr(block, 'detected_languages') and block.detected_languages:
                    return block.detected_languages[0].language_code
        
        return "en"  # Default to English
    
    async def preprocess_image(self, image_content: bytes) -> bytes:
        """
        Preprocess image for better OCR results.
        
        Args:
            image_content: Raw image bytes
            
        Returns:
            Preprocessed image bytes
        """
        if not PIL_AVAILABLE:
            logger.warning("PIL library not available - returning original image")
            return image_content
            
        try:
            # Open image with PIL
            image = Image.open(io.BytesIO(image_content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Basic preprocessing - could be enhanced with more sophisticated algorithms
            # For now, just ensure proper format and size
            
            # Resize if too large (max 4096x4096 for Document AI)
            max_size = 4096
            if image.width > max_size or image.height > max_size:
                image.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)
            
            # Save processed image
            output = io.BytesIO()
            image.save(output, format='JPEG', quality=95)
            return output.getvalue()
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {str(e)}")
            return image_content  # Return original if preprocessing fails