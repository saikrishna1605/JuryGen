"""
Document preprocessing utilities for improved OCR results.

This module provides:
- Image deskewing and denoising functions
- Document format detection and conversion
- Text extraction from various formats
- Image quality enhancement for OCR
"""

import asyncio
import logging
from typing import Tuple, Optional, Dict, Any, Union
from pathlib import Path
import io
import mimetypes

# Optional imports for document processing
try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    np = None
    NUMPY_AVAILABLE = False

try:
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps
    PIL_AVAILABLE = True
except ImportError:
    Image = ImageEnhance = ImageFilter = ImageOps = None
    PIL_AVAILABLE = False

try:
    import cv2
    CV2_AVAILABLE = True
except ImportError:
    cv2 = None
    CV2_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    magic = None
    MAGIC_AVAILABLE = False

from ..core.exceptions import DocumentFormatError, ProcessingException

logger = logging.getLogger(__name__)


class DocumentPreprocessor:
    """
    Document preprocessing utilities for enhanced OCR performance.
    
    Handles image enhancement, format detection, and document conversion
    to optimize text extraction accuracy.
    """
    
    def __init__(self):
        """Initialize document preprocessor."""
        if not PIL_AVAILABLE:
            logger.warning("PIL not available - image preprocessing disabled")
        if not CV2_AVAILABLE:
            logger.warning("OpenCV not available - advanced image processing disabled")
        if not PYPDF2_AVAILABLE:
            logger.warning("PyPDF2 not available - PDF text extraction disabled")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available - DOCX processing disabled")
        if not MAGIC_AVAILABLE:
            logger.warning("python-magic not available - file type detection disabled")
            
        self.supported_formats = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'application/msword': self._extract_doc_text,
            'text/plain': self._extract_plain_text,
            'image/jpeg': self._preprocess_image,
            'image/png': self._preprocess_image,
            'image/tiff': self._preprocess_image,
            'image/bmp': self._preprocess_image,
            'image/gif': self._preprocess_image,
        }
    
    async def detect_format(self, file_content: bytes, filename: str) -> str:
        """
        Detect document format from file content and filename.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            
        Returns:
            MIME type of the document
            
        Raises:
            DocumentFormatError: If format cannot be detected or is unsupported
        """
        try:
            # Try to detect MIME type from content using python-magic
            try:
                mime_type = magic.from_buffer(file_content, mime=True)
                logger.info(f"Detected MIME type from content: {mime_type}")
            except Exception as e:
                logger.warning(f"Magic detection failed: {str(e)}, falling back to filename")
                mime_type = None
            
            # Fallback to filename extension
            if not mime_type or mime_type == 'application/octet-stream':
                mime_type, _ = mimetypes.guess_type(filename)
                logger.info(f"Detected MIME type from filename: {mime_type}")
            
            # Validate detected format
            if not mime_type:
                raise DocumentFormatError(f"Cannot detect format for file: {filename}")
            
            if mime_type not in self.supported_formats:
                raise DocumentFormatError(f"Unsupported format: {mime_type}")
            
            return mime_type
            
        except Exception as e:
            logger.error(f"Format detection failed for {filename}: {str(e)}")
            raise DocumentFormatError(f"Format detection failed: {str(e)}") from e
    
    async def preprocess_document(
        self, 
        file_content: bytes, 
        filename: str,
        content_type: Optional[str] = None
    ) -> Tuple[bytes, Dict[str, Any]]:
        """
        Preprocess document for optimal OCR results.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: Optional MIME type (will be detected if not provided)
            
        Returns:
            Tuple of (processed_content, metadata)
            
        Raises:
            ProcessingException: If preprocessing fails
        """
        try:
            # Detect format if not provided
            if not content_type:
                content_type = await self.detect_format(file_content, filename)
            
            logger.info(f"Preprocessing {filename} as {content_type}")
            
            # Get appropriate preprocessor
            preprocessor = self.supported_formats.get(content_type)
            if not preprocessor:
                raise DocumentFormatError(f"No preprocessor for format: {content_type}")
            
            # Run preprocessing
            result = await preprocessor(file_content, filename)
            
            # Add metadata
            metadata = {
                'original_format': content_type,
                'original_size': len(file_content),
                'processed_size': len(result) if isinstance(result, bytes) else len(file_content),
                'preprocessing_applied': True
            }
            
            return result, metadata
            
        except Exception as e:
            logger.error(f"Document preprocessing failed for {filename}: {str(e)}")
            raise ProcessingException(f"Preprocessing failed: {str(e)}") from e
    
    async def _extract_pdf_text(self, file_content: bytes, filename: str) -> bytes:
        """Extract text from PDF or return original for OCR."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            # Try to extract text directly
            text_content = []
            has_extractable_text = False
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text.strip())
                        has_extractable_text = True
                    else:
                        text_content.append(f"[Page {page_num + 1} requires OCR]")
                except Exception as e:
                    logger.warning(f"Text extraction failed for page {page_num + 1}: {str(e)}")
                    text_content.append(f"[Page {page_num + 1} extraction failed]")
            
            # If we have extractable text, return it as plain text
            if has_extractable_text:
                full_text = "\n\n".join(text_content)
                return full_text.encode('utf-8')
            
            # Otherwise, return original PDF for OCR processing
            return file_content
            
        except Exception as e:
            logger.warning(f"PDF text extraction failed: {str(e)}, returning original for OCR")
            return file_content
    
    async def _extract_docx_text(self, file_content: bytes, filename: str) -> bytes:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract tables
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_cells.append(cell.text.strip())
                    if row_cells:
                        table_content.append(" | ".join(row_cells))
            
            # Combine content
            full_text = "\n\n".join(paragraphs)
            if table_content:
                full_text += "\n\nTables:\n" + "\n".join(table_content)
            
            return full_text.encode('utf-8')
            
        except Exception as e:
            raise ProcessingException(f"DOCX text extraction failed: {str(e)}") from e
    
    async def _extract_doc_text(self, file_content: bytes, filename: str) -> bytes:
        """Extract text from legacy DOC file."""
        # For legacy DOC files, we'd need additional libraries like python-docx2txt
        # For now, return original content and let OCR handle it
        logger.warning(f"Legacy DOC format not fully supported for {filename}, using OCR")
        return file_content
    
    async def _extract_plain_text(self, file_content: bytes, filename: str) -> bytes:
        """Process plain text file."""
        try:
            # Decode text and re-encode as UTF-8
            text = file_content.decode('utf-8', errors='replace')
            return text.encode('utf-8')
        except Exception as e:
            logger.warning(f"Text encoding issue for {filename}: {str(e)}")
            return file_content
    
    async def _preprocess_image(self, file_content: bytes, filename: str) -> bytes:
        """Preprocess image for better OCR results."""
        try:
            # Load image
            image = Image.open(io.BytesIO(file_content))
            
            # Convert to RGB if necessary
            if image.mode not in ['RGB', 'L']:
                image = image.convert('RGB')
            
            # Apply preprocessing pipeline
            processed_image = await self._enhance_image_for_ocr(image)
            
            # Save processed image
            output = io.BytesIO()
            processed_image.save(output, format='PNG', optimize=True)
            return output.getvalue()
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed for {filename}: {str(e)}")
            return file_content  # Return original if preprocessing fails
    
    async def _enhance_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Apply image enhancement techniques for better OCR.
        
        Args:
            image: PIL Image object
            
        Returns:
            Enhanced PIL Image
        """
        try:
            # Convert to numpy array for OpenCV operations
            img_array = np.array(image)
            
            # Apply deskewing if needed
            img_array = await self._deskew_image(img_array)
            
            # Apply denoising
            img_array = await self._denoise_image(img_array)
            
            # Convert back to PIL Image
            enhanced_image = Image.fromarray(img_array)
            
            # Apply PIL-based enhancements
            enhanced_image = await self._apply_pil_enhancements(enhanced_image)
            
            return enhanced_image
            
        except Exception as e:
            logger.warning(f"Image enhancement failed: {str(e)}")
            return image  # Return original if enhancement fails
    
    async def _deskew_image(self, img_array: np.ndarray) -> np.ndarray:
        """
        Detect and correct image skew.
        
        Args:
            img_array: Image as numpy array
            
        Returns:
            Deskewed image array
        """
        try:
            # Convert to grayscale if needed
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Apply edge detection
            edges = cv2.Canny(gray, 50, 150, apertureSize=3)
            
            # Detect lines using Hough transform
            lines = cv2.HoughLines(edges, 1, np.pi/180, threshold=100)
            
            if lines is not None:
                # Calculate skew angle
                angles = []
                for rho, theta in lines[:10]:  # Use first 10 lines
                    angle = theta * 180 / np.pi
                    # Convert to skew angle
                    if angle > 90:
                        angle = angle - 180
                    angles.append(angle)
                
                # Get median angle to avoid outliers
                if angles:
                    skew_angle = np.median(angles)
                    
                    # Only correct if skew is significant (> 0.5 degrees)
                    if abs(skew_angle) > 0.5:
                        logger.info(f"Correcting skew angle: {skew_angle:.2f} degrees")
                        
                        # Rotate image to correct skew
                        height, width = img_array.shape[:2]
                        center = (width // 2, height // 2)
                        rotation_matrix = cv2.getRotationMatrix2D(center, skew_angle, 1.0)
                        
                        # Calculate new dimensions to avoid cropping
                        cos_angle = abs(rotation_matrix[0, 0])
                        sin_angle = abs(rotation_matrix[0, 1])
                        new_width = int((height * sin_angle) + (width * cos_angle))
                        new_height = int((height * cos_angle) + (width * sin_angle))
                        
                        # Adjust translation
                        rotation_matrix[0, 2] += (new_width / 2) - center[0]
                        rotation_matrix[1, 2] += (new_height / 2) - center[1]
                        
                        # Apply rotation
                        if len(img_array.shape) == 3:
                            rotated = cv2.warpAffine(img_array, rotation_matrix, (new_width, new_height), 
                                                   borderValue=(255, 255, 255))
                        else:
                            rotated = cv2.warpAffine(img_array, rotation_matrix, (new_width, new_height), 
                                                   borderValue=255)
                        
                        return rotated
            
            return img_array
            
        except Exception as e:
            logger.warning(f"Deskewing failed: {str(e)}")
            return img_array
    
    async def _denoise_image(self, img_array: np.ndarray) -> np.ndarray:
        """
        Apply denoising to improve image quality.
        
        Args:
            img_array: Image as numpy array
            
        Returns:
            Denoised image array
        """
        try:
            # Apply different denoising based on image type
            if len(img_array.shape) == 3:
                # Color image - use Non-local Means Denoising
                denoised = cv2.fastNlMeansDenoisingColored(img_array, None, 10, 10, 7, 21)
            else:
                # Grayscale image
                denoised = cv2.fastNlMeansDenoising(img_array, None, 10, 7, 21)
            
            return denoised
            
        except Exception as e:
            logger.warning(f"Denoising failed: {str(e)}")
            return img_array
    
    async def _apply_pil_enhancements(self, image: Image.Image) -> Image.Image:
        """
        Apply PIL-based image enhancements.
        
        Args:
            image: PIL Image object
            
        Returns:
            Enhanced PIL Image
        """
        try:
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.2)  # Slight contrast boost
            
            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.1)  # Slight sharpness boost
            
            # Apply unsharp mask for better text clarity
            image = image.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))
            
            return image
            
        except Exception as e:
            logger.warning(f"PIL enhancement failed: {str(e)}")
            return image
    
    async def convert_to_searchable_pdf(
        self, 
        image_content: bytes, 
        ocr_text: str,
        filename: str
    ) -> bytes:
        """
        Convert image with OCR text to searchable PDF.
        
        Args:
            image_content: Original image bytes
            ocr_text: Extracted OCR text
            filename: Original filename
            
        Returns:
            Searchable PDF bytes
        """
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.utils import ImageReader
            
            # Create PDF buffer
            pdf_buffer = io.BytesIO()
            
            # Create PDF canvas
            c = canvas.Canvas(pdf_buffer, pagesize=letter)
            
            # Add image to PDF
            image = Image.open(io.BytesIO(image_content))
            img_reader = ImageReader(image)
            
            # Calculate image dimensions to fit page
            page_width, page_height = letter
            img_width, img_height = image.size
            
            # Scale image to fit page while maintaining aspect ratio
            scale = min(page_width / img_width, page_height / img_height)
            scaled_width = img_width * scale
            scaled_height = img_height * scale
            
            # Center image on page
            x = (page_width - scaled_width) / 2
            y = (page_height - scaled_height) / 2
            
            # Draw image
            c.drawImage(img_reader, x, y, scaled_width, scaled_height)
            
            # Add invisible text layer for searchability
            # This is a simplified implementation - a full solution would
            # position text based on OCR bounding boxes
            c.setFillColorRGB(1, 1, 1, alpha=0)  # Invisible text
            c.setFont("Helvetica", 8)
            
            # Add text in small chunks to avoid layout issues
            text_lines = ocr_text.split('\n')
            y_position = page_height - 50
            
            for line in text_lines[:50]:  # Limit to first 50 lines
                if y_position > 50:  # Don't go below page margin
                    c.drawString(50, y_position, line[:100])  # Limit line length
                    y_position -= 12
            
            # Save PDF
            c.save()
            
            return pdf_buffer.getvalue()
            
        except Exception as e:
            logger.warning(f"PDF conversion failed for {filename}: {str(e)}")
            # Return original image content if conversion fails
            return image_content
    
    def get_format_info(self, content_type: str) -> Dict[str, Any]:
        """
        Get information about a document format.
        
        Args:
            content_type: MIME type
            
        Returns:
            Format information dictionary
        """
        format_info = {
            'application/pdf': {
                'name': 'PDF',
                'supports_text_extraction': True,
                'requires_ocr': False,
                'preprocessing_available': True,
                'max_size_mb': 50
            },
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': {
                'name': 'DOCX',
                'supports_text_extraction': True,
                'requires_ocr': False,
                'preprocessing_available': False,
                'max_size_mb': 25
            },
            'application/msword': {
                'name': 'DOC',
                'supports_text_extraction': False,
                'requires_ocr': True,
                'preprocessing_available': False,
                'max_size_mb': 25
            },
            'text/plain': {
                'name': 'Text',
                'supports_text_extraction': True,
                'requires_ocr': False,
                'preprocessing_available': False,
                'max_size_mb': 10
            },
            'image/jpeg': {
                'name': 'JPEG',
                'supports_text_extraction': False,
                'requires_ocr': True,
                'preprocessing_available': True,
                'max_size_mb': 20
            },
            'image/png': {
                'name': 'PNG',
                'supports_text_extraction': False,
                'requires_ocr': True,
                'preprocessing_available': True,
                'max_size_mb': 20
            },
            'image/tiff': {
                'name': 'TIFF',
                'supports_text_extraction': False,
                'requires_ocr': True,
                'preprocessing_available': True,
                'max_size_mb': 30
            }
        }
        
        return format_info.get(content_type, {
            'name': 'Unknown',
            'supports_text_extraction': False,
            'requires_ocr': True,
            'preprocessing_available': False,
            'max_size_mb': 10
        })