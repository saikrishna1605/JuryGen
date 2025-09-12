"""
Document Service for handling document upload, storage, and retrieval.

This service handles:
- Document upload and validation
- Cloud Storage integration
- Document metadata management
- File type validation and processing
"""

import asyncio
import logging
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, timedelta
import uuid
import hashlib
import mimetypes
from pathlib import Path
import io

# Optional Google Cloud imports
try:
    from google.cloud import storage
    from google.api_core import exceptions as gcp_exceptions
    GOOGLE_CLOUD_AVAILABLE = True
except ImportError:
    storage = None
    gcp_exceptions = None
    GOOGLE_CLOUD_AVAILABLE = False

from ..core.config import get_settings
from ..core.exceptions import DocumentError, ValidationError, StorageError
from ..models.document import Document
from ..models.base import ProcessingStatus
from ..services.firestore import FirestoreService

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentService:
    """
    Service for managing document upload, storage, and retrieval.
    
    Handles file validation, cloud storage, and document metadata management.
    """
    
    def __init__(self):
        """Initialize the document service."""
        if not GOOGLE_CLOUD_AVAILABLE:
            logger.warning("Google Cloud libraries not available - document service functionality limited")
            self.storage_client = None
        else:
            try:
                self.storage_client = storage.Client()
            except Exception as e:
                logger.warning(f"Failed to initialize Storage client: {e}")
                self.storage_client = None
        
        try:
            self.firestore_service = FirestoreService()
        except Exception as e:
            logger.warning(f"Failed to initialize Firestore service: {e}")
            self.firestore_service = None
        
        # Configuration
        self.bucket_name = settings.STORAGE_BUCKET
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        
        # Supported file types
        self.supported_types = {
            "application/pdf": [".pdf"],
            "image/jpeg": [".jpg", ".jpeg"],
            "image/png": [".png"],
            "image/gif": [".gif"],
            "image/bmp": [".bmp"],
            "image/tiff": [".tiff", ".tif"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"],
            "application/msword": [".doc"]
        }
    
    async def upload_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: str,
        user_id: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Upload a document to cloud storage and save metadata.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type of the file
            user_id: User uploading the document
            metadata: Additional metadata
            
        Returns:
            Dictionary with upload result and document information
        """
        try:
            # Validate file
            self._validate_file(file_content, filename, content_type)
            
            # Generate document ID and storage path
            document_id = str(uuid.uuid4())
            file_extension = Path(filename).suffix.lower()
            storage_path = f"documents/{user_id}/{document_id}{file_extension}"
            
            # Upload to cloud storage
            storage_url = None
            if self.storage_client:
                try:
                    bucket = self.storage_client.bucket(self.bucket_name)
                    blob = bucket.blob(storage_path)
                    
                    # Upload file
                    blob.upload_from_string(
                        file_content,
                        content_type=content_type
                    )
                    
                    # Generate signed URL for access
                    storage_url = blob.generate_signed_url(
                        expiration=datetime.utcnow() + timedelta(hours=24),
                        method="GET"
                    )
                    
                    logger.info(f"Uploaded document to storage: {storage_path}")
                    
                except Exception as e:
                    logger.error(f"Storage upload failed: {e}")
                    raise StorageError(f"Failed to upload to storage: {e}")
            
            # Extract document metadata
            doc_metadata = self._extract_document_metadata(file_content, filename, content_type)
            
            # Create document record
            document_data = {
                "id": document_id,
                "user_id": user_id,
                "filename": filename,
                "content_type": content_type,
                "size_bytes": len(file_content),
                "storage_path": storage_path,
                "storage_url": storage_url,
                "status": ProcessingStatus.QUEUED.value,
                "metadata": {**doc_metadata, **(metadata or {})},
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow(),
                "hash": hashlib.sha256(file_content).hexdigest()
            }
            
            # Save to Firestore
            if self.firestore_service:
                await self.firestore_service.create_document(
                    collection="documents",
                    document_id=document_id,
                    data=document_data
                )
            
            return {
                "status": "success",
                "document_id": document_id,
                "filename": filename,
                "size_bytes": len(file_content),
                "storage_url": storage_url,
                "metadata": doc_metadata
            }
            
        except (ValidationError, StorageError):
            raise
        except Exception as e:
            logger.error(f"Document upload failed: {e}")
            raise DocumentError(f"Failed to upload document: {e}")
    
    async def get_document(self, document_id: str, user_id: str) -> Dict[str, Any]:
        """
        Retrieve document metadata and information.
        
        Args:
            document_id: Document identifier
            user_id: User requesting the document
            
        Returns:
            Document information
        """
        if not self.firestore_service:
            raise DocumentError("Firestore service not available")
            
        try:
            # Get document from Firestore
            document = await self.firestore_service.get_document(
                collection="documents",
                document_id=document_id
            )
            
            if not document:
                raise DocumentError("Document not found")
            
            # Check user access
            if document.get("user_id") != user_id:
                raise DocumentError("Access denied - document belongs to different user")
            
            return document
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Error retrieving document: {e}")
            raise DocumentError(f"Failed to retrieve document: {e}")
    
    async def list_user_documents(
        self,
        user_id: str,
        limit: int = 50,
        offset: int = 0
    ) -> List[Dict[str, Any]]:
        """
        List documents for a specific user.
        
        Args:
            user_id: User identifier
            limit: Maximum number of documents to return
            offset: Number of documents to skip
            
        Returns:
            List of document information
        """
        if not self.firestore_service:
            raise DocumentError("Firestore service not available")
            
        try:
            documents = await self.firestore_service.query_documents(
                collection="documents",
                filters=[("user_id", "==", user_id)],
                order_by=[("created_at", "desc")],
                limit=limit,
                offset=offset
            )
            
            return documents
            
        except Exception as e:
            logger.error(f"Error listing user documents: {e}")
            raise DocumentError(f"Failed to list documents: {e}")
    
    async def delete_document(self, document_id: str, user_id: str) -> Dict[str, Any]:
        """
        Delete a document and its associated files.
        
        Args:
            document_id: Document identifier
            user_id: User requesting deletion
            
        Returns:
            Deletion result
        """
        try:
            # Get document first to verify ownership
            document = await self.get_document(document_id, user_id)
            
            # Delete from cloud storage
            if self.storage_client and document.get("storage_path"):
                try:
                    bucket = self.storage_client.bucket(self.bucket_name)
                    blob = bucket.blob(document["storage_path"])
                    if blob.exists():
                        blob.delete()
                        logger.info(f"Deleted document from storage: {document['storage_path']}")
                except Exception as e:
                    logger.warning(f"Failed to delete from storage: {e}")
            
            # Delete from Firestore
            if self.firestore_service:
                await self.firestore_service.delete_document(
                    collection="documents",
                    document_id=document_id
                )
            
            return {"status": "success", "document_id": document_id}
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            raise DocumentError(f"Failed to delete document: {e}")
    
    async def update_document_status(
        self,
        document_id: str,
        status: ProcessingStatus,
        metadata: Optional[Dict[str, Any]] = None
    ) -> None:
        """
        Update document processing status.
        
        Args:
            document_id: Document identifier
            status: New status
            metadata: Additional metadata to update
        """
        if not self.firestore_service:
            raise DocumentError("Firestore service not available")
            
        try:
            update_data = {
                "status": status.value,
                "updated_at": datetime.utcnow()
            }
            
            if metadata:
                update_data["metadata"] = metadata
            
            await self.firestore_service.update_document(
                collection="documents",
                document_id=document_id,
                data=update_data
            )
            
        except Exception as e:
            logger.error(f"Error updating document status: {e}")
            raise DocumentError(f"Failed to update document status: {e}")
    
    async def get_signed_download_url(
        self,
        document_id: str,
        user_id: str,
        expiration_hours: int = 1
    ) -> str:
        """
        Generate a signed URL for document download.
        
        Args:
            document_id: Document identifier
            user_id: User requesting the URL
            expiration_hours: URL expiration time in hours
            
        Returns:
            Signed download URL
        """
        if not self.storage_client:
            raise DocumentError("Cloud Storage not available")
            
        try:
            # Get document to verify ownership and get storage path
            document = await self.get_document(document_id, user_id)
            
            if not document.get("storage_path"):
                raise DocumentError("Document not found in storage")
            
            # Generate signed URL
            bucket = self.storage_client.bucket(self.bucket_name)
            blob = bucket.blob(document["storage_path"])
            
            signed_url = blob.generate_signed_url(
                expiration=datetime.utcnow() + timedelta(hours=expiration_hours),
                method="GET"
            )
            
            return signed_url
            
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Error generating signed URL: {e}")
            raise DocumentError(f"Failed to generate download URL: {e}")
    
    def _validate_file(self, file_content: bytes, filename: str, content_type: str) -> None:
        """
        Validate uploaded file.
        
        Args:
            file_content: File content bytes
            filename: Original filename
            content_type: MIME type
            
        Raises:
            ValidationError: If file is invalid
        """
        # Check file size
        if len(file_content) > self.max_file_size:
            raise ValidationError(f"File too large. Maximum size is {self.max_file_size // (1024*1024)}MB")
        
        # Check file type
        if content_type not in self.supported_types:
            raise ValidationError(f"Unsupported file type: {content_type}")
        
        # Check file extension
        file_extension = Path(filename).suffix.lower()
        if file_extension not in self.supported_types[content_type]:
            raise ValidationError(f"File extension {file_extension} doesn't match content type {content_type}")
        
        # Validate file content
        self._validate_document_content(file_content, filename, content_type)
    
    def _validate_document_content(self, file_content: bytes, filename: str, content_type: str) -> bool:
        """
        Validate document content matches declared type.
        
        Args:
            file_content: File content bytes
            filename: Original filename
            content_type: Declared MIME type
            
        Returns:
            True if valid
            
        Raises:
            ValidationError: If content is invalid
        """
        if content_type == "application/pdf":
            # Check PDF header
            if not file_content.startswith(b'%PDF-'):
                raise ValidationError("Invalid PDF file - missing PDF header")
        
        elif content_type.startswith("image/"):
            # Basic image validation
            try:
                from PIL import Image
                image = Image.open(io.BytesIO(file_content))
                image.verify()
            except ImportError:
                logger.warning("PIL not available - skipping image validation")
            except Exception:
                raise ValidationError(f"Invalid image file: {filename}")
        
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            # Basic DOCX/DOC validation
            if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # DOCX files are ZIP archives
                if not file_content.startswith(b'PK'):
                    raise ValidationError("Invalid DOCX file - not a valid ZIP archive")
        
        return True
    
    def _extract_document_metadata(
        self,
        file_content: bytes,
        filename: str,
        content_type: str
    ) -> Dict[str, Any]:
        """
        Extract metadata from document content.
        
        Args:
            file_content: File content bytes
            filename: Original filename
            content_type: MIME type
            
        Returns:
            Dictionary with extracted metadata
        """
        metadata = {
            "size": len(file_content),
            "creation_date": datetime.utcnow().isoformat(),
            "file_extension": Path(filename).suffix.lower()
        }
        
        try:
            if content_type == "application/pdf":
                # Extract PDF metadata
                try:
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                except ImportError:
                    logger.warning("PyPDF2 not available - skipping PDF metadata extraction")
                    return metadata
                metadata["page_count"] = len(pdf_reader.pages)
                
                # Try to get PDF info
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    if pdf_info.get("/Title"):
                        metadata["title"] = str(pdf_info["/Title"])
                    if pdf_info.get("/Author"):
                        metadata["author"] = str(pdf_info["/Author"])
                    if pdf_info.get("/CreationDate"):
                        metadata["pdf_creation_date"] = str(pdf_info["/CreationDate"])
            
            elif content_type.startswith("image/"):
                # Extract image metadata
                try:
                    from PIL import Image
                    from PIL.ExifTags import TAGS
                    
                    image = Image.open(io.BytesIO(file_content))
                    metadata["image_width"] = image.width
                    metadata["image_height"] = image.height
                    metadata["image_mode"] = image.mode
                    
                    # Extract EXIF data if available
                    if hasattr(image, '_getexif') and image._getexif():
                        exif_data = {}
                        for tag_id, value in image._getexif().items():
                            tag = TAGS.get(tag_id, tag_id)
                            exif_data[tag] = str(value)
                        metadata["exif"] = exif_data
                        
                except Exception as e:
                    logger.warning(f"Failed to extract image metadata: {e}")
            
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ]:
                # Extract DOCX metadata
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(io.BytesIO(file_content))
                    
                    metadata["paragraph_count"] = len(doc.paragraphs)
                    metadata["table_count"] = len(doc.tables)
                    
                    # Core properties
                    if doc.core_properties.title:
                        metadata["title"] = doc.core_properties.title
                    if doc.core_properties.author:
                        metadata["author"] = doc.core_properties.author
                    if doc.core_properties.created:
                        metadata["docx_creation_date"] = doc.core_properties.created.isoformat()
                        
                except Exception as e:
                    logger.warning(f"Failed to extract DOCX metadata: {e}")
        
        except Exception as e:
            logger.warning(f"Failed to extract document metadata: {e}")
        
        return metadata


# Singleton instance
document_service = DocumentService()